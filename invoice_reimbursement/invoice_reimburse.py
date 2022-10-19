# coding=utf-8
"""
This script is the prototype of invoice reimbursement automation for everbright trust
"""
import os
import math

from random import randint, sample, choice
from docx import Document
from docx2pdf import convert

from .models import CompanyReim, ClientReim, FieldReim


class InvoiceReimburse:
    def __init__(self, reim_input):
        self.invoice_info = reim_input
        self.invoice_info['amount'] = float(self.invoice_info['amount'])

        #  set the multiple of average dining fee base based on dining type
        self.base_multiplpe_dict = {
            '工作招待': 1,
            '一般招待': 2,
            '商务招待': 3,
        }
        self.base_multiplpe = self.base_multiplpe_dict[self.invoice_info['dining_type']]

        #  initialize class attributes
        self.clients_str = ''  # readable clients string
        self.guests = 0  # number of guests
        self.avai_com_objs = []  # a list of CompanyReim objects, the company in which has more clients than self.guests
        self.table = ''  #  the table instance generated using docx.Document.tables in the word attachment


    def attachment_gen(self):
        """
        generate the attachment file required
        """
        base_dir = os.path.dirname(os.path.abspath(__file__))
        print(base_dir)
        file_name = 'template【发票报销】.docx'
        template_path = base_dir + '/filesources/' + file_name
        doc = Document(template_path)
        #  use docx to open an attachment template. There is a table inside. Populate the table with generated info.
        self.table = doc.tables[0]
        self.cell_manuver(0, 1, self.field)
        self.cell_manuver(1, 1, self.invoice_info['date'])
        self.cell_manuver(1, 3, self.confirmed_company)
        self.cell_manuver(2, 1, self.invoice_info['restaurant'])
        self.cell_manuver(3, 3, self.clients_str)
        self.cell_manuver(4, 1, self.guests_str)
        self.cell_manuver(4, 3, self.hosts_str)
        self.cell_manuver(5, 1, self.amount_str)
        self.cell_manuver(5, 3, self.ave_amount_str)
        attach_name_docx = self.invoice_info['date'] + '-' + self.amount_str + '.docx'
        attach_path_docx = base_dir + '/filedownload/' + attach_name_docx
        doc.save(attach_path_docx)
        # conver docx to pdf as attachment
        attach_name = self.invoice_info['date'] + '-' + self.amount_str + '.pdf'
        attach_path = base_dir + '/filedownload/' + attach_name
        convert(attach_path_docx, attach_path)
        os.remove(attach_path_docx)
        return attach_name

    def cell_manuver(self, x, y, txt):
        """
        populate the cell with coordinates of (x,y) with text of txt
        """
        cell = self.table.cell(x, y)
        runs = cell.paragraphs[0].runs
        runs[0].clear()
        cell.paragraphs[0].add_run(text=txt)

    def who_to_dine(self):
        """
        automatically select the guests and generate other relevant info for reimbusement
        return a dict for page display of reimbursement info
        """
        self.diners_num()
        self.avai_com()

        random_com_obj = choice(self.avai_com_objs)  #  avai_com_objs is a list of objects, thus using choice here
        self.confirmed_company = random_com_obj.company
        client_objs = [client_obj for client_obj in ClientReim.objects.filter(company=random_com_obj)]

        random_client_objs = sample(client_objs, self.guests)
        # print(random_client_objs)

        # create a readable client string, for example, "投资部门+投资经理+王+总，"。

        for client_obj in random_client_objs:
            self.clients_str += client_obj.department + client_obj.position + client_obj.client + '总' + '，'
        self.clients_str = self.clients_str[:-1]  # remove the last comma

        last = FieldReim.objects.filter(company=random_com_obj).count() - 1
        idx = randint(0, last)
        random_field_obj = FieldReim.objects.filter(company=random_com_obj)[idx]
        # print(random_field_obj)

        #  texts prep for page display and attachment table population
        self.field = "就" + random_field_obj.field + "进行洽谈"
        self.ave_amount = self.invoice_info['amount'] / self.diners
        self.amount_str = '{:,.2f}'.format(self.invoice_info['amount']) + "元"
        self.ave_amount_str = '{:,.2f}'.format(self.ave_amount) + "元"
        self.hosts_str = str(self.hosts) + "人"
        self.guests_str = str(self.guests) + "人"
        """
        print("________________________")
        print("招待单位：" + self.confirmed_company)
        print("招待事宜：" + self.field)
        print("招待人员：" + self.clients_str)
        print("接待人数：" + self.hosts_str + "；招待人数：" + self.guests_str)
        print("总金额 ：" +  self.amount_str)
        print("人均金额：" + self.ave_amount_str)
        print("________________________")
        """

        return {
            'confirmed_company': self.confirmed_company,
            'field': self.field,
            'clients': self.clients_str,
            'hosts': self.hosts_str,
            'guests': self.guests_str,
            'amount': self.amount_str,
            'ave_amount': self.ave_amount_str,
        }


    def diners_num(self):
        """
        take in preset number of diners_num if there is any, and generate the number of diners for the hosts and guests
        """
        if self.invoice_info['preset_number']:
            # print('存在设定的用餐人数')
            self.diners = int(self.invoice_info['preset_number'])
        else:
            min_number = max(2, math.ceil(self.invoice_info['amount'] / (200*self.base_multiplpe)))
            max_number = min(max(2, int(self.invoice_info['amount'] // (140*self.base_multiplpe))), 6)
            self.diners = randint(min_number, max_number)

            """
            print('不存在设定的用餐人数，自动产生用餐人数')
            print('最小人数：' + str(min_number))
            print('最大人数：' + str(max_number))
            """

        # print('总人数：' + str(self.diners))
        self.hosts = randint(1, int(self.diners / 2))
        self.guests = self.diners - self.hosts
        # print('其中接待人数设置为: ' + str(self.hosts))

    def avai_com(self):
        """
        return available companies (companies that have more clients than guests)
        """
        # ('self.avai_com starts')
        coms_objs = CompanyReim.objects.all()
        for com_obj in coms_objs:
            #  what if add a field to CompanyReim to get the client number of the company
            clients_objs = ClientReim.objects.filter(company=com_obj)
            if len(clients_objs) >= self.guests:
                self.avai_com_objs.append(com_obj)
        # print(self.avai_com_objs)
        """
        for company, clients_in_a_company in self.clients_dict.items():
            clients_number_in_a_company = len(clients_in_a_company)
            if clients_number_in_a_company >= self.guests:
                avai_com.append(company)
        """

    """
    def clients_dict_gen(self):

        self.clients_dict = {}
        all_companies = CompanyReim.objects.all()

        for row in tuple(self.ws.values)[1:]:
            name, depar, position, *field = row[1:]
            field = list(filter(None, field))
            depar = depar or ''  # only depar could be vacant
            if row[0] not in self.clients_dict.keys():
                self.clients_dict[row[0]] = []
            self.clients_dict[row[0]].append([name, depar, position, field])
    """

if __name__ == '__main__':
    reim_input = {
        'amount': 3000,
        'restaurant': '好吃的餐厅',
        'date': '2022-10-07',
        'preset_number': None,
        'dining_type': '商务招待',
    }
    ir = InvoiceReimburse(reim_input)
    ir.who_to_dine()
    # ir.attachment_gen()




